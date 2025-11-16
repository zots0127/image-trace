import io
import json
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path
import zipfile
import tempfile
import os

from PIL import Image
import fitz  # PyMuPDF
from docx import Document
from pptx import Presentation
import pdf2image


class DocumentProcessor:
    """文档处理器，用于从各种文档格式中提取图片"""

    SUPPORTED_FORMATS = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/vnd.openxmlformats-officedocument.presentationml.presentation': 'pptx',
        'application/msword': 'doc',
        'application/vnd.ms-powerpoint': 'ppt'
    }

    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()

    def __del__(self):
        # 清理临时文件
        try:
            import shutil
            shutil.rmtree(self.temp_dir, ignore_errors=True)
        except:
            pass

    def get_document_format(self, filename: str, content_type: str) -> Optional[str]:
        """根据文件名和MIME类型确定文档格式"""
        if content_type in self.SUPPORTED_FORMATS:
            return self.SUPPORTED_FORMATS[content_type]

        # 根据文件扩展名判断
        ext = Path(filename).suffix.lower()
        format_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.pptx': 'pptx',
            '.ppt': 'ppt'
        }
        return format_map.get(ext)

    def extract_metadata(self, file_data: bytes, doc_format: str) -> Dict[str, Any]:
        """提取文档元数据"""
        metadata = {
            'format': doc_format,
            'file_size': len(file_data)
        }

        try:
            if doc_format == 'pdf':
                with fitz.open(stream=file_data, filetype="pdf") as doc:
                    metadata.update({
                        'page_count': doc.page_count,
                        'title': doc.metadata.get('title', ''),
                        'author': doc.metadata.get('author', ''),
                        'subject': doc.metadata.get('subject', ''),
                        'creator': doc.metadata.get('creator', ''),
                        'producer': doc.metadata.get('producer', ''),
                        'creation_date': doc.metadata.get('creationDate', ''),
                        'modification_date': doc.metadata.get('modDate', '')
                    })

            elif doc_format == 'docx':
                doc = Document(io.BytesIO(file_data))
                metadata.update({
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables)
                })

            elif doc_format == 'pptx':
                prs = Presentation(io.BytesIO(file_data))
                metadata.update({
                    'slide_count': len(prs.slides)
                })
        except Exception as e:
            metadata['metadata_error'] = str(e)

        return metadata

    def extract_images_from_pdf(self, file_data: bytes) -> List[Tuple[bytes, Dict[str, Any]]]:
        """从PDF中提取图片"""
        images = []

        try:
            # 方法1: 使用PyMuPDF提取嵌入图片
            with fitz.open(stream=file_data, filetype="pdf") as doc:
                img_count = 0

                for page_num in range(len(doc)):
                    page = doc[page_num]

                    # 提取页面上的图片
                    image_list = page.get_images(full=True)

                    for img_index, img in enumerate(image_list):
                        try:
                            # 获取图片引用
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)

                            # 跳过CMYK图像（转换复杂）
                            if pix.n - pix.alpha < 4:
                                img_data = pix.tobytes("png")

                                extraction_metadata = {
                                    'source_page': page_num + 1,
                                    'image_index': img_index + 1,
                                    'extraction_method': 'pymupdf_embedded',
                                    'width': pix.width,
                                    'height': pix.height,
                                    'color_space': pix.colorspace.name if pix.colorspace else 'unknown'
                                }

                                images.append((img_data, extraction_metadata))
                                img_count += 1

                            pix = None
                        except Exception as e:
                            print(f"Failed to extract image {img_index} from page {page_num}: {e}")
                            continue

                # 方法2: 如果没有找到嵌入图片，将每页渲染为图片
                if not images:
                    try:
                        # 使用pdf2image将PDF页面转换为图像
                        pil_images = pdf2image.convert_from_bytes(
                            file_data,
                            dpi=200,
                            output_folder=None,
                            fmt='jpeg'
                        )

                        for page_num, pil_image in enumerate(pil_images):
                            img_buffer = io.BytesIO()
                            pil_image.save(img_buffer, format='JPEG', quality=85)
                            img_data = img_buffer.getvalue()

                            extraction_metadata = {
                                'source_page': page_num + 1,
                                'image_index': 1,
                                'extraction_method': 'pdf2image_render',
                                'width': pil_image.width,
                                'height': pil_image.height,
                                'color_space': 'RGB'
                            }

                            images.append((img_data, extraction_metadata))

                    except Exception as e:
                        print(f"Failed to render PDF pages as images: {e}")

        except Exception as e:
            print(f"Failed to process PDF: {e}")

        return images

    def extract_images_from_docx(self, file_data: bytes) -> List[Tuple[bytes, Dict[str, Any]]]:
        """从DOCX中提取图片"""
        images = []

        try:
            # DOCX文件实际上是一个ZIP文件
            with zipfile.ZipFile(io.BytesIO(file_data)) as docx_zip:
                # 查找media文件夹中的图片
                media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]

                for img_index, media_file in enumerate(media_files):
                    try:
                        with docx_zip.open(media_file) as img_file:
                            img_data = img_file.read()

                            extraction_metadata = {
                                'source_file': media_file,
                                'image_index': img_index + 1,
                                'extraction_method': 'docx_media_zip',
                                'file_size': len(img_data)
                            }

                            images.append((img_data, extraction_metadata))

                    except Exception as e:
                        print(f"Failed to extract image {media_file}: {e}")
                        continue

        except Exception as e:
            print(f"Failed to process DOCX: {e}")

        return images

    def extract_images_from_pptx(self, file_data: bytes) -> List[Tuple[bytes, Dict[str, Any]]]:
        """从PPTX中提取图片"""
        images = []

        try:
            # PPTX文件也是一个ZIP文件
            with zipfile.ZipFile(io.BytesIO(file_data)) as pptx_zip:
                # 查找media文件夹中的图片
                media_files = [f for f in pptx_zip.namelist() if f.startswith('ppt/media/')]

                for img_index, media_file in enumerate(media_files):
                    try:
                        with pptx_zip.open(media_file) as img_file:
                            img_data = img_file.read()

                            extraction_metadata = {
                                'source_file': media_file,
                                'image_index': img_index + 1,
                                'extraction_method': 'pptx_media_zip',
                                'file_size': len(img_data)
                            }

                            images.append((img_data, extraction_metadata))

                    except Exception as e:
                        print(f"Failed to extract image {media_file}: {e}")
                        continue

        except Exception as e:
            print(f"Failed to process PPTX: {e}")

        return images

    def extract_images(self, file_data: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """
        从文档中提取图片的主入口函数

        Args:
            file_data: 文档文件数据
            filename: 文件名
            content_type: MIME类型

        Returns:
            包含提取结果的字典
        """
        doc_format = self.get_document_format(filename, content_type)
        if not doc_format:
            return {
                'success': False,
                'error': 'Unsupported document format',
                'images': []
            }

        # 提取文档元数据
        metadata = self.extract_metadata(file_data, doc_format)

        # 根据格式提取图片
        images = []
        try:
            if doc_format == 'pdf':
                images = self.extract_images_from_pdf(file_data)
            elif doc_format == 'docx':
                images = self.extract_images_from_docx(file_data)
            elif doc_format == 'pptx':
                images = self.extract_images_from_pptx(file_data)
            else:
                return {
                    'success': False,
                    'error': f'Format {doc_format} not yet supported for image extraction',
                    'images': []
                }

        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to extract images: {str(e)}',
                'images': []
            }

        # 转换图片格式（如果需要）
        processed_images = []
        for img_data, extraction_metadata in images:
            try:
                # 尝试转换为JPEG格式以节省空间
                with Image.open(io.BytesIO(img_data)) as img:
                    # 如果是PNG且有透明度，转换为白色背景JPEG
                    if img.mode in ('RGBA', 'LA'):
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'RGBA':
                            background.paste(img, mask=img.split()[-1])
                        else:
                            background.paste(img, mask=img.split()[-1])
                        img = background

                    # 转换为JPEG
                    img_buffer = io.BytesIO()
                    img.save(img_buffer, format='JPEG', quality=85, optimize=True)
                    processed_data = img_buffer.getvalue()

                    # 更新元数据
                    extraction_metadata.update({
                        'processed_size': len(processed_data),
                        'original_size': len(img_data),
                        'final_format': 'JPEG'
                    })

                    processed_images.append((processed_data, extraction_metadata))

            except Exception as e:
                # 如果转换失败，保留原始图片
                extraction_metadata['processing_error'] = str(e)
                processed_images.append((img_data, extraction_metadata))

        return {
            'success': True,
            'format': doc_format,
            'metadata': metadata,
            'image_count': len(processed_images),
            'images': processed_images
        }


# 全局文档处理器实例
document_processor = DocumentProcessor()